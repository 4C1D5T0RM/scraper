import requests
from bs4 import BeautifulSoup
from docx import Document

# ---------- TEXT CLEANING ----------
def get_text_with_links(el):
    """Return text from an element including hyperlink text."""
    parts = []
    for part in el.strings:
        text = part.strip()
        if not text:
            continue
        if hasattr(part.parent, "name") and part.parent.name == "a":
            href = part.parent.get("href", "").strip()
            if href:
                parts.append(f"{text} ({href})")
            else:
                parts.append(text)
        else:
            parts.append(text)
    return " ".join(parts)

# ---------- TXT EXPORT ----------
def process_list_items_txt(container_name, element, level, output_lines):
    lis = element.find_all('li')
    for li in lis:
        text = get_text_with_links(li)
        indent = '  ' * level
        if text:
            output_lines.append(f"{indent}{container_name} | LI (level {level}): {text}")
        for sublist in li.find_all(['ul', 'ol']):
            output_lines.append(f"{indent}{container_name} | {sublist.name.upper()}:")
            process_list_items_txt(container_name, sublist, level + 1, output_lines)

def extract_with_nesting_txt(container_name, container, output_lines, level=0):
    if not container:
        return
    elements = container.find_all(['h1','h2','h3','h4','h5','p','ul','ol'])
    for el in elements:
        tag = el.name.upper()
        text = get_text_with_links(el)
        indent = '  ' * level
        if not text:
            continue
        if tag.startswith('H'):
            heading_level = int(tag[1])
            output_lines.append(f"{indent}{container_name} | {tag} (level {heading_level}): {text}")
        elif tag == 'P':
            output_lines.append(f"{indent}{container_name} | P: {text}")
        elif tag in ['UL', 'OL']:
            output_lines.append(f"{indent}{container_name} | {tag}:")
            process_list_items_txt(container_name, el, level + 1, output_lines)

def save_as_txt(filename, soup, scrape_header, scrape_body, scrape_footer):
    output_lines = []
    header = soup.find('header')
    footer = soup.find('footer')
    body = soup.find('body')

    header_texts = set()
    footer_texts = set()

    if scrape_header and header:
        output_lines.append("HEADER:")
        header_texts = collect_text_set(header)
        extract_with_nesting_txt("HEADER", header, output_lines)
    else:
        if header:
            header.extract()

    if scrape_body and body:
        output_lines.append("BODY:")
        extract_with_nesting_txt_skip_set("BODY", body, output_lines, header_texts, footer_texts)

    if scrape_footer and footer:
        output_lines.append("FOOTER:")
        footer_texts = collect_text_set(footer)
        extract_with_nesting_txt("FOOTER", footer, output_lines)

    with open(f"{filename}.txt", "w", encoding="utf-8") as f:
        for line in output_lines:
            f.write(line + "\n")
    print(f"✅ Saved as {filename}.txt")

def collect_text_set(container):
    texts = set()
    for el in container.find_all(['h1','h2','h3','h4','h5','p','li']):
        text = get_text_with_links(el)
        if text:
            texts.add(text)
    return texts

def extract_with_nesting_txt_skip_set(container_name, container, output_lines, skip_header, skip_footer, level=0):
    if not container:
        return
    elements = container.find_all(['h1','h2','h3','h4','h5','p','ul','ol'])
    for el in elements:
        tag = el.name.upper()
        text = get_text_with_links(el)
        indent = '  ' * level
        if not text or text in skip_header or text in skip_footer:
            continue
        if tag.startswith('H'):
            heading_level = int(tag[1])
            output_lines.append(f"{indent}{container_name} | {tag} (level {heading_level}): {text}")
        elif tag == 'P':
            output_lines.append(f"{indent}{container_name} | P: {text}")
        elif tag in ['UL', 'OL']:
            output_lines.append(f"{indent}{container_name} | {tag}:")
            process_list_items_txt(container_name, el, level + 1, output_lines)

# ---------- DOCX EXPORT ----------
def add_runs_from_element(el, para):
    """Add runs to a paragraph including text from hyperlinks."""
    added = False
    for part in el.strings:
        text = part.strip()
        if not text:
            continue
        if hasattr(part.parent, "name") and part.parent.name == "a":
            href = part.parent.get("href", "").strip()
            if href:
                para.add_run(f"{text} ({href})")
            else:
                para.add_run(text)
        else:
            para.add_run(text)
        added = True
    return added

def process_list_items_docx(element, level, doc):
    lis = element.find_all('li')
    for li in lis:
        para = doc.add_paragraph()
        para.style = 'List Paragraph'
        para.paragraph_format.left_indent = level * 300
        added = add_runs_from_element(li, para)
        if not added:
            doc._body._element.remove(para._element)
        for sublist in li.find_all(['ul','ol']):
            process_list_items_docx(sublist, level + 1, doc)

def extract_with_nesting_docx(container, doc, level=0):
    if not container:
        return
    elements = container.find_all(['h1','h2','h3','h4','h5','p','ul','ol'])
    for el in elements:
        tag = el.name.upper()
        indent = level * 300
        if tag.startswith('H'):
            heading_level = int(tag[1])
            para = doc.add_paragraph()
            para.style = f'Heading {heading_level}'
            para.paragraph_format.left_indent = indent
            added = add_runs_from_element(el, para)
            if not added:
                doc._body._element.remove(para._element)
        elif tag == 'P':
            para = doc.add_paragraph()
            para.style = 'Normal'
            para.paragraph_format.left_indent = indent
            added = add_runs_from_element(el, para)
            if not added:
                doc._body._element.remove(para._element)
        elif tag in ['UL', 'OL']:
            process_list_items_docx(el, level + 1, doc)

def save_as_docx(filename, soup, scrape_header, scrape_body, scrape_footer):
    doc = Document()
    header = soup.find('header')
    footer = soup.find('footer')
    body = soup.find('body')

    if scrape_header and header:
        doc.add_paragraph("HEADER:")
        extract_with_nesting_docx(header, doc)
    else:
        if header:
            header.extract()

    if scrape_body and body:
        doc.add_paragraph("BODY:")
        extract_with_nesting_docx(body, doc)

    if scrape_footer and footer:
        doc.add_paragraph("FOOTER:")
        extract_with_nesting_docx(footer, doc)

    doc.save(f"{filename}.docx")
    print(f"✅ Saved as {filename}.docx")

# ---------- MAIN ----------
def scrape_website(url):
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        scrape_header = input("🔍 Scrape HEADER? (yes/no): ").strip().lower() in ['yes', 'y']
        scrape_body = input("🔍 Scrape BODY? (yes/no): ").strip().lower() in ['yes', 'y']
        scrape_footer = input("🔍 Scrape FOOTER? (yes/no): ").strip().lower() in ['yes', 'y']

        filename = input("📄 Filename (no extension): ").strip() or "scraped_content"
        filetype = input("💾 Save as txt or docx? ").strip().lower()

        if filetype == 'txt':
            save_as_txt(filename, soup, scrape_header, scrape_body, scrape_footer)
        else:
            save_as_docx(filename, soup, scrape_header, scrape_body, scrape_footer)

        other_format = 'docx' if filetype == 'txt' else 'txt'
        if input(f"💾 Export to {other_format} too? (yes/no): ").strip().lower() in ['yes', 'y']:
            alt_filename = input(f"📄 {other_format} filename (no extension): ").strip() or f"{filename}_alt"
            if other_format == 'txt':
                save_as_txt(alt_filename, soup, scrape_header, scrape_body, scrape_footer)
            else:
                save_as_docx(alt_filename, soup, scrape_header, scrape_body, scrape_footer)

    except requests.exceptions.RequestException as e:
        print(f"❌ Error: {e}")

def main():
    while True:
        url = input("🌐 URL to scrape (or type 'exit'): ").strip()
        if url.lower() == 'exit':
            print("👋 Exiting.")
            break
        if not url.startswith('http'):
            print("⚠ Please include http/https in the URL.")
            continue
        scrape_website(url)

if __name__ == "__main__":
    main()
